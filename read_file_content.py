# 标准库导入
import re
import logging

from pathlib import Path
from typing import Optional, Union, List
# 第三方库导入
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment

import os
import platform
import subprocess
import tempfile
import time
import shutil # 用于查找可执行文件和删除目录


# 文件操作

# 配置日志
logging.basicConfig(
    level=logging.WARN,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
# 显式设置root logger级别
logging.getLogger().setLevel(logging.WARN)

def read_file_content(file_path: Union[str, Path]) -> Optional[str]:
    """读取文件内容并返回去除首尾空格的字符串
    
    Args:
        file_path: 文件路径，支持字符串或Path对象
        
    Returns:
        str: 文件内容（去除首尾空格）
        None: 文件不存在或读取失败
        
    Raises:
        FileNotFoundError: 文件不存在
        IOError: 文件读取失败
    """
    path = Path(file_path) if isinstance(file_path, str) else file_path
    
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    try:
        with path.open('r', encoding='utf-8') as file:
            return file.read().strip()
    except UnicodeDecodeError as e:
        raise IOError(f"文件解码失败: {path}") from e
    except Exception as e:
        raise IOError(f"读取文件出错: {path}") from e


# 常量定义
EXCEL_COLUMN_NAMES = {
    "requirement": 2,  # 需求列索引
    "process": 4       # 流程列索引
}

DEFAULT_FONT_STYLES = {
    "heading1": {"name": "黑体", "size": Pt(22), "color": RGBColor(0, 0, 0)},
    "heading2": {"name": "黑体", "size": Pt(16), "color": RGBColor(0, 0, 0)},
    "body": {"name": "宋体", "size": Pt(10.5)}
}

def save_content_to_file(
    file_name: str, 
    output_dir: Union[str, Path],
    content: str, 
    content_type: str = "text"
) -> Path:
    """
    保存内容到文件，支持不同类型的文件。

    Args:
        file_name: 输出文件名（不含扩展名）
        output_dir: 输出目录路径
        content: 要保存的字符串内容
        content_type: 内容类型 ("text", "json", "markdown", "other")
    """

    try:
        # 1. 提取基础文件名
        base_name = os.path.splitext(file_name)[0]

        # 2. 创建输出目录（如果不存在）
        os.makedirs(output_dir, exist_ok=True)

        # 4. 根据 content_type 确定文件扩展名和保存逻辑
        if content_type == "json":
            output_filename = os.path.join(output_dir, f"{base_name}.json")
            with open(output_filename, "w", encoding="utf-8") as f:
                f.write(content)
        elif content_type == "text":
            output_filename = os.path.join(output_dir, f"{base_name}.txt")
            with open(output_filename, "w", encoding="utf-8") as f:
                f.write(content)

        elif content_type == "markdown":
            output_filename = os.path.join(output_dir, f"{base_name}.md")
            with open(output_filename, "w", encoding="utf-8") as f:
                f.write(content)
        elif content_type == 'xlsx':
            output_filename = os.path.join(output_dir, f"{base_name}.xlsx")
            df = markdown_table_to_df(content)
            if df is not None:
                df.to_excel(output_filename, index=False)  # 保存为 Excel

            # 合并表格单元格
            merge_cells_by_column(output_filename, "Sheet1")
        elif content_type == 'docx':
            ##先读取excel 文件
            excel_file = os.path.join(output_dir, f"{base_name}.xlsx")
            output_filename = os.path.join(output_dir, f"{base_name}.docx")
            create_function_design_doc(excel_file, output_filename)
        else:  # 默认情况，可以根据需要添加更多类型
            # 获取原始文件的扩展名
            original_extension = os.path.splitext(file_name)[1]
            output_filename = os.path.join(output_dir, f"{base_name}{original_extension}")
            with open(output_filename, "w", encoding="utf-8") as f:
                f.write(content)

        logging.info(f"已创建文件: {output_filename}")

    except Exception as e:
        logging.error(f"处理文件 {file_name} 时发生错误", exc_info=True)


def extract_content_from_requst(text, extract_type: str = "total_rows"):
    """
      Args:
        text: 包含数字的文本字符串。

      Returns:
        提取到的数字 (整数)，如果没有找到，则返回 None。
      """
    if extract_type == 'total_rows' :
        match = re.search(r"表格总行数要求：(\d+)", text)  #
        return int(match.group(1))
    if extract_type == 'request_name':
        match = re.search(r"客户需求：(.*)", text)  # 需求名称
        return match.group(1)
    else:
        return None


def markdown_table_to_df(table_text):
    # 分割行
    lines = table_text.strip().split('\n')

    # --- 表头处理 (关键改进) ---
    header_line = lines[0]
    separator_line = lines[1]

    # 获取表头单元格数量（基于分隔符行）
    header_cols = [s.strip() for s in re.split(r'(?<!\\)\|', separator_line)[1:-1]]
    num_cols = len(header_cols)

    # 解析表头行，并根据分隔符行的数量进行调整
    header = [h.strip() for h in re.split(r'(?<!\\)\|', header_line)[1:-1]]
    header = (header + [''] * (num_cols - len(header)))[:num_cols]  # 补齐或截断

    # --- 数据行处理 ---
    data = []
    for line in lines[2:]:
        row = [cell.strip() for cell in re.split(r'(?<!\\)\|', line)[1:-1]]
        row = (row + [''] * (num_cols - len(row)))  # 补齐
        row = [cell.replace("<br>", "\n") for cell in row]
        data.append(row)

    df = pd.DataFrame(data, columns=header)
    return df


def process_markdown_table(markdown_table_string, num_cols_to_process=5):
    """
    处理Markdown表格，合并指定数量列中连续相同内容的单元格。

    Args:
        markdown_table_string: Markdown表格的字符串。
        num_cols_to_process: 要处理的列数。

    Returns:
        处理后的Markdown表格字符串。
    """

    lines = markdown_table_string.strip().split('\n')
    header = lines[0:2]
    data_rows = lines[2:]

    table_data = [re.split(r'\s*\|\s*', row.strip('|')) for row in data_rows]

    num_cols = len(table_data[0])

    # 限制处理的列数
    num_cols_to_process = min(num_cols_to_process, num_cols)

    for col_index in range(num_cols_to_process):
        previous_value = None
        for row_index in range(len(table_data)):
            current_value = table_data[row_index][col_index]
            if current_value == previous_value:
                table_data[row_index][col_index] = ""
            else:
                previous_value = current_value

    processed_rows = ['| ' + ' | '.join(row) + ' |' for row in table_data]
    result_table = '\n'.join(header + processed_rows)

    return result_table



def apply_font_style(run, style_name: str) -> None:
    """应用预定义的字体样式
    
    Args:
        run: docx的Run对象
        style_name: 样式名称（heading1/heading2/body）
    """
    style = DEFAULT_FONT_STYLES[style_name]
    run.font.name = style["name"]
    run.font.size = style["size"]
    if "color" in style:
        run.font.color.rgb = style["color"]

def create_function_design_doc(excel_file: Union[str, Path], docx_file: Union[str, Path]) -> None:
    """
    根据COSMIC Excel表格生成功能设计文档（.docx）。
    (处理合并单元格，设置标题颜色为黑色)

    Args:
        excel_file: COSMIC Excel文件的路径。
        docx_file:  生成的Word文档的路径。
    """

    workbook = load_workbook(filename=excel_file)
    sheet = workbook.active

    document = Document()

    # 第4章 系统功能设计 (固定内容)
    heading = document.add_heading("第4章 系统功能设计", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.runs[0]
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色

    section_num = 4
    current_requirement = None

    for row in sheet.iter_rows(min_row=2, values_only=True):
        requirement = row[EXCEL_COLUMN_NAMES["requirement"]]
        process = row[EXCEL_COLUMN_NAMES["process"]]

        if requirement:
            current_requirement = requirement
            section_num += 0.1
            section_num = round(section_num, 1)

            # 添加功能用户需求标题
            heading = document.add_heading(f"{section_num} {current_requirement}", level=2)
            apply_font_style(heading.runs[0], "heading2")
            current_process_num = 1  # 重置流程计数器

        if process:
            paragraph = document.add_paragraph(f"（{current_process_num}）{process}")
            run = paragraph.runs[0]
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            current_process_num += 1

    document.save(docx_file)


def merge_cells_by_column(filename, sheetname):
    """
    按列合并 Excel 单元格。

    Args:
        filename: Excel 文件名。
        sheetname: 要处理的 Sheet 名称。
    """
    workbook = load_workbook(filename)
    sheet = workbook[sheetname]

    # 处理 A 到 E 列 (0-4)
    for col_index in range(5):  # 列索引从0开始
        start_row = None
        start_value = None
        end_row = None  # 新增：记录批次的结束行

        # 循环每一行，从第二行开始（跳过标题行）
        for row_index in range(2, sheet.max_row + 1):
            cell = sheet.cell(row=row_index, column=col_index + 1)
            cell_value = cell.value

            if cell_value is not None and cell_value != "":  # 非空单元格
                if start_row is None:  # 找到第一个非空单元格
                    start_row = row_index
                    start_value = cell_value
                    end_row = row_index  # 初始时，结束行就是起始行
                elif cell_value != start_value:  # 遇到不同内容的非空单元格
                    # 合并单元格 (如果 end_row 有效)
                    if end_row is not None:
                        sheet.merge_cells(
                            start_row=start_row,
                            start_column=col_index + 1,
                            end_row=end_row,  # 使用 end_row
                            end_column=col_index + 1,
                        )
                        merged_cell = sheet.cell(row=start_row, column=col_index + 1)
                        merged_cell.alignment = Alignment(
                            horizontal="center", vertical="center", wrap_text=True
                        )
                        merged_cell.value = start_value

                    # 重置起始位置和结束位置
                    start_row = row_index
                    start_value = cell_value
                    end_row = row_index
                else:  # 遇到相同内容的非空单元格
                    end_row = row_index  # 更新结束行

            elif start_row is not None:  # 遇到空单元格，且已经有起始位置
                end_row = row_index      # 更新结束行

        # 处理最后一批单元格（如果存在）
        if start_row is not None and end_row is not None:
            sheet.merge_cells(
                start_row=start_row,
                start_column=col_index + 1,
                end_row=end_row,  # 使用 end_row
                end_column=col_index + 1,
            )
            merged_cell = sheet.cell(row=start_row, column=col_index + 1)
            merged_cell.alignment = Alignment(
                horizontal="center", vertical="center", wrap_text=True
            )
            merged_cell.value = start_value

    workbook.save(filename)  # 保存修改


def merge_temp_files(temp_files: List[Path]) -> str:
    """合并临时Markdown表格文件"""
    full_content = []

    for i, file_path in enumerate(sorted(temp_files)):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().splitlines()

            if i == 0:
                # 保留第一个文件的完整头
                full_content.extend(content)
            else:
                # 跳过后续文件的头两行（标题和分隔符）
                full_content.extend(content[2:])

    return "\n".join(full_content)


def process_text(content):
    """
    数据清洗
    删除文本中“业务流程（必填）” 后的内容，并删除特殊字符“”。

    Args:
        content: 包含文本内容的字符串。

    Returns:
        处理后的字符串。
    """

    markers = [
        "能力复用评估",
        "业务流程图/时序图（如涉及，必填）"
    ]
    # 特殊字符 U+FFFD (REPLACEMENT CHARACTER)
    # 在 Python 字符串中可以直接使用，或者用 Unicode 转义 \uFFFD
    special_char = "\x01"
    # 查找所有标记的所有出现位置
    occurrences = []
    for marker in markers:
        start_pos = 0
        while True:
            index = content.find(marker, start_pos)
            if index == -1:
                break  # 当前标记在此后的文本中未找到
            occurrences.append(index)  # 记录找到的位置
            # 移动搜索起始点到当前找到的标记之后，避免重复查找同一位置
            start_pos = index + 1  # 加1即可，find会找到第一个字符匹配的位置

    # 对所有找到的位置进行排序
    occurrences.sort()

    # 初始化 content_to_process 为原始内容
    content_to_process = content
    truncation_point = -1  # -1 表示不截断

    # 检查是否有至少两次出现（任意标记组合）
    if len(occurrences) >= 2:
        # 第二次出现的位置是排序后列表的第二个元素 (索引为 1)
        truncation_point = occurrences[1]
        logging.debug(f"在位置 {truncation_point} 找到第 2 个标记（来自列表 {markers}）。")
        content_to_process = content[:truncation_point]
    else:
        logging.warning(f"在文件  中找到的标记（来自列表 {markers}）总数少于 2 个。将不执行截断。")
        # content_to_process 保持为原始内容

    # 删除特殊字符 (无论是否截断，都执行此操作)
    processed_content = content_to_process.replace(special_char, "")
    processed_content = processed_content.replace('\x07', "")
    while processed_content.find('\r\r') > -1:
        processed_content = processed_content.replace('\r\r', "\r")

        # --- 步骤 3: 根据关键词删除行 ---
    lines = processed_content.splitlines(keepends=True)
    
    # 查找第一个"客户需求规格说明书"和第二个"背景描述"的位置
    start_indices = [i for i, line in enumerate(lines) if "客户需求规格说明书" in line]
    end_indices = [i for i, line in enumerate(lines) if "背景描述" in line]
    
    if len(start_indices) > 0 and len(end_indices) > 1:
        start_line = start_indices[0]
        end_line = end_indices[1]  # 取第二个"背景描述"
        
        if start_line < end_line:
            # 保留开始行之前和结束行之后的内容
            final_lines = lines[:start_line] + lines[end_line:]
            processed_content = "".join(final_lines)
            logging.debug(f"已删除从第{start_line+1}行到第{end_line}行的内容")
        else:
            logging.warning("'客户需求规格说明书'出现在第二个'背景描述'之后，不执行删除")
    else:
        logging.warning("未找到足够的匹配行，不执行删除操作")




    return processed_content

# --- Linux .doc 转 .docx (使用 LibreOffice) ---
def _check_libreoffice():
    """检查 libreoffice 命令是否存在"""
    return shutil.which('libreoffice') is not None

def _convert_doc_to_docx_linux(doc_path: str, output_dir: str) -> str | None:
    """
    在 Linux 上使用 LibreOffice 将 .doc 转换为 .docx。

    参数:
        doc_path: 输入的 .doc 文件路径。
        output_dir: 用于存放转换后文件的目录。

    返回:
        成功时返回转换后的 .docx 文件完整路径，失败时返回 None。
    """
    if not _check_libreoffice():
        logging.error("在 Linux 上转换 .doc 文件需要安装 LibreOffice。")
        logging.error("请运行: sudo apt-get update && sudo apt-get install libreoffice")
        return None

    abs_doc_path = os.path.abspath(doc_path)
    abs_output_dir = os.path.abspath(output_dir)

    if not os.path.exists(abs_doc_path):
        logging.error(f"输入文件未找到: {abs_doc_path}")
        return None
    if not os.path.isdir(abs_output_dir):
         logging.error(f"输出目录不存在: {abs_output_dir}")
         return None

    logging.debug(f"Linux: 使用 LibreOffice 转换: {abs_doc_path} -> {abs_output_dir}")
    try:
        # 构建 LibreOffice 命令
        # 使用明确的过滤器通常更可靠
        # 你可以通过 soffice --help 查看可用的过滤器
        # "MS Word 2007 XML" 是常见的 .docx 过滤器名称
        command = [
            'libreoffice',
            '--headless', # 无头模式，不显示 GUI
            '--convert-to', 'docx:"MS Word 2007 XML"', # 指定输出格式和过滤器
            '--outdir', abs_output_dir, # 指定输出目录
            abs_doc_path # 输入文件
        ]
        logging.debug(f"执行命令: {' '.join(command)}")

        # 执行转换命令，设置超时（例如 60 秒）
        result = subprocess.run(command, capture_output=True, text=True, timeout=60, check=False) # check=False 手动检查

        # 检查命令执行结果
        if result.returncode != 0:
            logging.error(f"LibreOffice 转换失败。返回码: {result.returncode}")
            logging.error(f"错误信息 (stderr):\n{result.stderr}")
            logging.error(f"输出信息 (stdout):\n{result.stdout}")
            return None

        # 构造预期的输出文件名
        base_name = os.path.splitext(os.path.basename(abs_doc_path))[0]
        expected_docx_path = os.path.join(abs_output_dir, base_name + ".docx")

        # 检查输出文件是否存在
        if os.path.exists(expected_docx_path):
            logging.debug(f"LibreOffice 转换成功: {expected_docx_path}")
            return expected_docx_path
        else:
            logging.error(f"LibreOffice 命令执行成功，但未找到预期的输出文件: {expected_docx_path}")
            logging.error(f"请检查输出目录 '{abs_output_dir}' 的内容和权限。")
            logging.error(f"LibreOffice stdout:\n{result.stdout}")
            logging.error(f"LibreOffice stderr:\n{result.stderr}")
            return None

    except subprocess.TimeoutExpired:
        logging.error(f"LibreOffice 转换超时 ({60} 秒)。文件可能过大或 LibreOffice 进程卡住。")
        # (可选) 尝试杀死可能残留的 soffice 进程
        # import psutil
        # for proc in psutil.process_iter(['pid', 'name']):
        #     if 'soffice' in proc.info['name'].lower():
        #         print(f"尝试终止残留的 LibreOffice 进程: PID {proc.info['pid']}")
        #         try:
        #             proc.terminate()
        #             proc.wait(timeout=5) # 等待进程结束
        #         except (psutil.NoSuchProcess, psutil.TimeoutExpired, psutil.AccessDenied):
        #             pass # 忽略错误
        return None
    except FileNotFoundError:
         logging.error("'libreoffice' 命令未找到。请确保 LibreOffice 已安装并添加到系统 PATH。")
         return None
    except Exception as e:
        logging.error(f"Linux 转换过程中发生未知错误: {e}")
        return None

# --- Windows .doc 转 .docx (使用 win32com) ---
def _convert_doc_to_docx_windows(doc_path: str, docx_path: str) -> str | None:
    """
    在 Windows 上使用 Word COM 将 .doc 转换为 .docx。

    参数:
        doc_path: 输入的 .doc 文件路径。
        docx_path: 输出的 .docx 文件路径。

    返回:
        成功时返回转换后的 .docx 文件路径，失败时返回 None。
    """
    # --- Windows 特定导入 ---
    if platform.system() == "Windows":
        try:
            import win32com.client as win32
            import pythoncom
            # Word 文件格式常量 (wdSaveFormat)
            WD_FORMAT_DOCX = 16  # wdFormatXMLDocument
        except ImportError:
            logging.warning("Windows 平台需要 'pywin32' 库。请运行 'pip install pywin32'")
            win32 = None  # 标记为不可用
    else:
        win32 = None  # 在非 Windows 平台上标记为不可用

    if not win32:
        logging.error("Windows 转换需要 'pywin32' 库且 Word 已安装。")
        return None

    word = None
    doc = None
    abs_doc_path = os.path.abspath(doc_path)
    abs_docx_path = os.path.abspath(docx_path)

    if not os.path.exists(abs_doc_path):
        logging.error(f"输入文件未找到: {abs_doc_path}")
        return None

    logging.debug(f"Windows: 使用 Word COM 转换: {abs_doc_path} -> {abs_docx_path}")

    try:
        pythoncom.CoInitialize() # 初始化 COM 环境
        # 尝试清除缓存可能遇到的问题
        # 注意：EnsureDispatch 在缓存损坏时可能依然失败，需要手动删 gen_py
        try:
             word = win32.gencache.EnsureDispatch('Word.Application')
        except AttributeError as e:
             if 'CLSIDToClassMap' in str(e):
                 logging.error("检测到 win32com 缓存问题。尝试清理缓存...")
                 # 定位并删除 gen_py 目录
                 gencache_path = os.path.join(os.path.dirname(pythoncom.__file__), '..', 'win32com', 'gen_py')
                 if os.path.exists(gencache_path):
                     logging.error(f"删除缓存目录: {gencache_path}")
                     shutil.rmtree(gencache_path, ignore_errors=True)
                     # 重新尝试 Dispatch
                     logging.error("缓存已清理，重试 Dispatch...")
                     word = win32.Dispatch('Word.Application') # 使用 Dispatch 避免再次触发 EnsureDispatch 的缓存生成
                 else:
                     logging.error(f"缓存目录未找到: {gencache_path}")
                     raise e # 如果找不到缓存目录，重新抛出原始错误
             else:
                 raise e # 抛出其他 AttributeError
        except Exception as e:
            logging.error(f"初始化 Word Application 时出错: {e}")
            raise # 重新抛出未能处理的异常

        word.Visible = False
        word.DisplayAlerts = False

        logging.info(f"正在打开文档: {abs_doc_path}")
        doc = word.Documents.Open(abs_doc_path, ReadOnly=True)

        logging.info(f"正在另存为 .docx 格式: {abs_docx_path}")
        doc.SaveAs(abs_docx_path, FileFormat=WD_FORMAT_DOCX)

        logging.info("Word COM 转换成功！")
        return abs_docx_path

    except pythoncom.com_error as e:
        logging.error(f"Windows 转换过程中发生 COM 错误: {e}")
        # (hresult, strerror, excepinfo, argerror) = e.args
        # print(f"HRESULT: {hresult}, Error: {strerror}")
        # if excepinfo: print(f"Exception Info: {excepinfo}")
        return None
    except Exception as e:
        logging.error(f"Windows 转换过程中发生未知错误: {e}")
        return None
    finally:
        # 确保资源被释放
        if doc:
            try:
                doc.Close(False)
                logging.debug("文档已关闭。")
            except pythoncom.com_error as e:
                logging.error(f"关闭文档时发生 COM 错误: {e}")
        if word:
            try:
                word.Quit()
                logging.debug("Word 应用已退出。")
            except pythoncom.com_error as e:
                logging.error(f"退出 Word 时发生 COM 错误: {e}")
        # 仅在初始化成功后才反初始化
        if 'pythoncom' in locals() and hasattr(pythoncom, 'CoUninitialize'):
             pythoncom.CoUninitialize()

# --- 主读取函数 ---
def read_word_document(file_path: str) -> str:
    """
    跨平台读取 Word 文档内容 (.doc 或 .docx)。
    对于 .doc 文件，会尝试将其转换为 .docx 后再读取。

    参数:
        file_path: Word 文件路径。

    返回:
        文档文本内容。

    依赖:
        - 所有平台: python-docx
        - Windows: pywin32, Microsoft Word
        - Linux: libreoffice
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件未找到: {file_path}")

    original_path = os.path.abspath(file_path)
    file_ext = os.path.splitext(original_path)[1].lower()
    docx_to_read = None
    temp_dir = None # 用于存放临时转换文件和目录
    temp_docx_path = None # 临时 .docx 文件路径

    try:
        if file_ext == '.docx':
            logging.info(f"直接读取 .docx 文件: {original_path}")
            docx_to_read = original_path
        elif file_ext == '.doc':
            logging.info(f"检测到 .doc 文件，需要转换: {original_path}")
            # 创建一个临时目录来存放转换后的文件
            temp_dir = tempfile.mkdtemp(prefix="wordconv_")
            logging.info(f"创建临时目录: {temp_dir}")
            # 构造临时的 .docx 文件名
            base_name = os.path.splitext(os.path.basename(original_path))[0]
            temp_docx_filename = f"{base_name}_{int(time.time())}.docx"

            current_os = platform.system()

            if current_os == "Windows":
                temp_docx_path = os.path.join(temp_dir, temp_docx_filename)
                converted_path = _convert_doc_to_docx_windows(original_path, temp_docx_path)
                if converted_path and os.path.exists(converted_path):
                    docx_to_read = converted_path
                else:
                    raise RuntimeError(f"Windows: 无法将 .doc 转换为 .docx: {original_path}")

            elif current_os == "Linux":
                # Linux 转换函数需要输出目录
                converted_path = _convert_doc_to_docx_linux(original_path, temp_dir)
                if converted_path and os.path.exists(converted_path):
                    docx_to_read = converted_path
                    # 注意：LibreOffice 可能不会使用我们指定的临时文件名，
                    # 所以我们使用它实际返回的路径。
                    temp_docx_path = converted_path # 记录实际的临时文件路径以便删除
                else:
                    raise RuntimeError(f"Linux: 无法将 .doc 转换为 .docx: {original_path}")

            else:
                raise OSError(f"不支持的操作系统: {current_os}. 无法处理 .doc 文件。")

        else:
            raise ValueError(f"不支持的文件格式: {file_ext}. 只支持 .doc 和 .docx")

        # --- 现在统一读取 .docx 文件 ---
        if not docx_to_read:
             raise RuntimeError("未能确定要读取的 .docx 文件路径。")

        logging.info(f"开始读取 .docx 内容: {docx_to_read}")
        content = ""
        try:
            doc = Document(docx_to_read)
            # 提取所有段落的文本
            # 你也可以提取表格内容等，根据需要调整
            paragraphs = [p.text for p in doc.paragraphs]
            # 简单地用换行符连接段落
            content = '\n'.join(paragraphs)
            logging.info(".docx 文件内容读取成功。")
        except Exception as e:
            raise RuntimeError(f"读取 .docx 文件 '{docx_to_read}' 时失败: {e}") from e

        content = process_text(content)
        return content
    finally:
        # --- 清理临时文件和目录 ---
        # 优先删除具体的临时 docx 文件（如果路径已知且存在）
        if temp_docx_path and os.path.exists(temp_docx_path):
             try:
                 os.remove(temp_docx_path)
                 logging.info(f"已删除临时 .docx 文件: {temp_docx_path}")
             except OSError as e:
                 logging.warning(f"删除临时文件失败: {temp_docx_path}, 错误: {e}")
        # 然后删除整个临时目录（如果创建了）
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                logging.info(f"已删除临时目录: {temp_dir}")
            except OSError as e:
                logging.warning(f"删除临时目录失败: {temp_dir}, 错误: {e}")
