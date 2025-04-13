import json
from docxtpl import DocxTemplate
import docx
import os
from docx.shared import Cm
import logging
from typing import Optional
import datetime # 导入 datetime 模块

# 使用与 requirement_analysis 相同的日志记录器或配置
logger = logging.getLogger(__name__) # 或者 logging.getLogger('requirement_analysis')

def replace_placeholder_with_image(doc_path, placeholder_text, image_path, width_cm=None):
    """
    在 Word 文档中查找占位符文本并将其替换为图片。
    (函数体保持不变，但添加日志记录)
    """
    if not os.path.exists(doc_path):
        logger.error(f"错误：找不到文档 '{doc_path}'")
        return False
    if not os.path.exists(image_path):
        logger.error(f"错误：找不到图片 '{image_path}'")
        return False

    try:
        doc = docx.Document(doc_path)
        placeholder_found = False

        # --- 遍历文档中的段落 ---
        for paragraph in doc.paragraphs:
            if placeholder_text in paragraph.text:
                logger.debug(f"在段落中找到占位符: '{placeholder_text}' in '{doc_path}'")
                inline = paragraph.runs
                for i in range(len(inline) - 1, -1, -1):
                    if placeholder_text in inline[i].text:
                        inline[i].text = inline[i].text.replace(placeholder_text, '')
                paragraph.clear()
                run = paragraph.add_run()
                if width_cm:
                    run.add_picture(image_path, width=Cm(width_cm))
                else:
                    run.add_picture(image_path)
                placeholder_found = True
                logger.debug(f"已在段落插入图片 '{os.path.basename(image_path)}'")
                break # 假设占位符在文档中是唯一的

        # --- （可选）遍历表格中的单元格 ---
        if not placeholder_found:
             for table in doc.tables:
                 for row in table.rows:
                     for cell in row.cells:
                         for paragraph in cell.paragraphs:
                             if placeholder_text in paragraph.text:
                                 logger.debug(f"在表格单元格中找到占位符: '{placeholder_text}' in '{doc_path}'")
                                 inline = paragraph.runs
                                 for i in range(len(inline) - 1, -1, -1):
                                     if placeholder_text in inline[i].text:
                                         inline[i].text = inline[i].text.replace(placeholder_text, '')
                                 paragraph.clear()
                                 run = paragraph.add_run()
                                 if width_cm:
                                     run.add_picture(image_path, width=Cm(width_cm))
                                 else:
                                     run.add_picture(image_path)
                                 placeholder_found = True
                                 logger.debug(f"已在表格单元格插入图片 '{os.path.basename(image_path)}'")
                                 # break out of inner loops if found
                                 break
                         if placeholder_found: break
                     if placeholder_found: break
                 if placeholder_found: break

        if placeholder_found:
            doc.save(doc_path)
            logger.info(f"图片 '{os.path.basename(image_path)}' 已成功插入到 '{doc_path}' 的 '{placeholder_text}' 位置。")
            return True
        else:
            logger.warning(f"警告：在文档 '{doc_path}' 中未找到占位符 '{placeholder_text}'")
            return False

    except Exception as e:
        logger.error(f"在文档 '{doc_path}' 中替换占位符 '{placeholder_text}' 为图片 '{image_path}' 时发生错误: {e}", exc_info=True)
        return False


def generate_word_document(requirement_name:str, json_data_path: str, template_path: str, output_doc_path: str, word_text: Optional[str] = None, image_path: Optional[str] = None, image_placeholder: str = 'sequence_diagram_mermaid', image_width_cm: Optional[float] = None):
    """
    使用 JSON 数据和模板生成 Word 文档，并可选地插入图片和文本。

    Args:
        requirement_name (str): 需求名称 (用于填充模板)。
        json_data_path (str): 输入的 JSON 数据文件路径。
        template_path (str): Word 模板文件路径 (.docx)。
        output_doc_path (str): 输出的 Word 文档文件路径。
        word_text (Optional[str]): 要填充到 {{word_text}} 占位符的文本内容 (例如，合并的 Markdown)。
        template_path (str): Word 模板文件路径 (.docx)。
        output_doc_path (str): 输出的 Word 文档文件路径。
        image_path (Optional[str]): 要插入的图片文件路径。如果为 None，则不插入图片。
        image_placeholder (str): Word 模板中用于图片插入的占位符文本。
        image_width_cm (Optional[float]): 插入图片的宽度（厘米）。
    Returns:
        bool: 操作是否成功。
    """
    logger.info(f"开始生成 Word 文档: {output_doc_path}")
    logger.debug(f"  JSON 数据源: {json_data_path}")
    logger.debug(f"  模板文件: {template_path}")
    if image_path:
        logger.debug(f"  图片文件: {image_path}")
        logger.debug(f"  图片占位符: {image_placeholder}")

    # 1. 检查输入文件是否存在
    if not os.path.exists(json_data_path):
        logger.error(f"错误：找不到 JSON 数据文件 '{json_data_path}'")
        return False
    if not os.path.exists(template_path):
        logger.error(f"错误：找不到 Word 模板文件 '{template_path}'")
        return False
    if image_path and not os.path.exists(image_path):
        logger.warning(f"警告：找不到指定的图片文件 '{image_path}'，将不执行图片插入。")
        image_path = None # 设为 None 以跳过插入

    try:
        # 2. 加载 JSON 数据
        with open(json_data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # 3. 加载 Word 模板
        doc = DocxTemplate(template_path)

        # 4. 准备上下文数据
        context = {
            'requirement_description': data.get('requirement_description', {}),
            'system_status': data.get('system_status', {}),
            'functional_architecture_diagram': data.get('functional_architecture_diagram', {}),
            'overall_description': data.get('requirement_description', {}).get('overall_description', ''),
            'construction_goals': data.get('requirement_description', {}).get('construction_goals', []),
            'necessity': data.get('requirement_description', {}).get('necessity', []),
            'system_overview': data.get('system_status', {}).get('system_overview', []),
            'implemented_features': data.get('system_status', {}).get('implemented_features', []),
            'existing_problems': data.get('system_status', {}).get('existing_problems', ''),
            # 添加当前月份
            'MONTH': datetime.datetime.now().strftime('%m'),
            # 添加需求名称和功能点文本
            'req_name': requirement_name,
            'word_text': word_text or '', # 如果 word_text 为 None，则使用空字符串
        }
        logger.debug("上下文数据准备完成。")

        # 5. 渲染模板
        doc.render(context)
        logger.debug("Word 模板渲染完成。")

        # 6. 保存渲染后的文档
        # 确保输出目录存在
        output_dir = os.path.dirname(output_doc_path)
        if output_dir: # 如果路径包含目录
             os.makedirs(output_dir, exist_ok=True)
        doc.save(output_doc_path)
        logger.info(f"基于模板的 Word 文档已生成: {output_doc_path}")

        # 7. 插入图片（如果提供了图片路径且文件存在）
        if image_path:
            logger.info(f"尝试将图片 '{image_path}' 插入到 '{output_doc_path}'...")
            success = replace_placeholder_with_image(output_doc_path, image_placeholder, image_path, width_cm=image_width_cm)
            if success:
                logger.info("图片插入成功。")
            # else: # replace_placeholder_with_image 内部已有日志记录
            #     logger.warning("图片插入未成功（可能未找到占位符或发生错误）。")
        else:
            logger.info("未提供有效图片路径，跳过图片插入步骤。")

        return True

    except FileNotFoundError as fnf_err:
        logger.error(f"生成 Word 文档时文件未找到: {fnf_err}")
        return False
    except json.JSONDecodeError as json_err:
        logger.error(f"解析 JSON 文件 '{json_data_path}' 失败: {json_err}")
        return False
    except Exception as e:
        logger.error(f"生成 Word 文档 '{output_doc_path}' 时发生未知错误: {e}", exc_info=True)
        return False

# 注意：移除了原脚本末尾的独立执行代码
# if __name__ == "__main__":
#    ...
